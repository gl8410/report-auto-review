import os
import io
import uuid
import json
from typing import List, Optional
from fastapi import UploadFile, HTTPException
from sqlmodel import Session, select
from backend.core.config import settings
from backend.core.db import engine
from backend.models.document import Document, DocumentStatus
from backend.models.chunk import DocumentChunk
from backend.integrations.vector_store import ingest_chunks_to_chroma, delete_document_from_chroma, dynamic_chunk_text

async def extract_text_from_file(content: bytes, filename: str) -> str:
    """
    Extract text content from different file formats
    """
    filename_lower = filename.lower()

    if filename_lower.endswith('.pdf'):
        # Extract text from PDF
        try:
            from pypdf import PdfReader
            pdf_reader = PdfReader(io.BytesIO(content))
            text_parts = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n".join(text_parts)
        except Exception as e:
            print(f"PDF extraction error: {e}")
            raise ValueError(f"Failed to extract text from PDF: {e}")

    elif filename_lower.endswith('.docx'):
        # Extract text from DOCX
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            return "\n".join(text_parts)
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {e}")

    else:
        # Plain text files (.txt, .md)
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            return content.decode('gbk', errors='ignore')

async def process_document_background(doc_id: str, file_content: bytes, filename: str):
    """Background task to parse and vectorize document."""
    from sqlmodel import Session as SyncSession

    with SyncSession(engine) as session:
        doc = session.get(Document, doc_id)
        if not doc:
            print(f"Document {doc_id} not found for processing")
            return

        try:
            # Update status to PARSING
            doc.status = DocumentStatus.PARSING.value
            session.add(doc)
            session.commit()

            # Extract text from document
            text_content = await extract_text_from_file(file_content, filename)

            if not text_content or len(text_content.strip()) < 10:
                raise ValueError("Document contains no extractable text")

            # Store text in meta_info for reference
            doc.meta_info = json.dumps({
                "text_length": len(text_content),
                "filename": filename
            })
            session.add(doc)
            session.commit()

            # Generate chunks
            chunks_data = dynamic_chunk_text(text_content)
            
            # Save chunks to SQL
            for chunk in chunks_data:
                db_chunk = DocumentChunk(
                    document_id=doc_id,
                    chunk_index=chunk["index"],
                    content=chunk["text"],
                    word_count=chunk["word_count"],
                    sentence_count=chunk["sentence_count"]
                )
                session.add(db_chunk)
            session.commit()

            # Ingest to ChromaDB
            chunk_count = await ingest_chunks_to_chroma(doc_id, chunks_data, filename)

            # Update status to INDEXED
            doc.status = DocumentStatus.INDEXED.value
            doc.meta_info = json.dumps({
                "text_length": len(text_content),
                "chunk_count": chunk_count,
                "filename": filename
            })
            session.add(doc)
            session.commit()

            print(f"Successfully indexed document {filename} with {chunk_count} chunks")

        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"Error processing document {filename}: {repr(e)}")
            doc.status = DocumentStatus.FAILED.value
            doc.meta_info = json.dumps({"error": str(e)})
            session.add(doc)
            session.commit()

class DocumentService:
    @staticmethod
    def get_documents(session: Session) -> List[Document]:
        """Get all documents ordered by upload time (newest first)."""
        return session.exec(select(Document).order_by(Document.upload_time.desc())).all()

    @staticmethod
    def get_document(session: Session, doc_id: str) -> Optional[Document]:
        """Get a specific document."""
        return session.get(Document, doc_id)

    @staticmethod
    def get_document_chunks(session: Session, doc_id: str) -> List[DocumentChunk]:
        """Get all chunks for a document."""
        return session.exec(select(DocumentChunk).where(DocumentChunk.document_id == doc_id).order_by(DocumentChunk.chunk_index)).all()

    @staticmethod
    async def upload_document(
        session: Session,
        file: UploadFile,
        background_tasks
    ) -> Document:
        """Upload a document for review (supports .pdf, .docx, .doc)."""
        filename = file.filename or "unknown"

        # Validate file type
        valid_extensions = ['.pdf', '.docx', '.doc']
        ext = os.path.splitext(filename)[1].lower()
        if ext not in valid_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type. Allowed: {', '.join(valid_extensions)}"
            )

        # Read file content
        file_content = await file.read()

        # Ensure uploads directory exists
        os.makedirs(settings.UPLOADS_DIR, exist_ok=True)

        # Generate unique filename for storage
        doc_id = str(uuid.uuid4())
        storage_filename = f"{doc_id}{ext}"
        storage_path = os.path.join(settings.UPLOADS_DIR, storage_filename)

        # Save file to disk
        with open(storage_path, 'wb') as f:
            f.write(file_content)

        # Create document record
        doc = Document(
            id=doc_id,
            filename=filename,
            storage_path=storage_path,
            status=DocumentStatus.UPLOADED.value
        )
        session.add(doc)
        session.commit()
        session.refresh(doc)

        # Start background processing
        background_tasks.add_task(process_document_background, doc_id, file_content, filename)

        return doc

    @staticmethod
    async def delete_document(session: Session, doc_id: str) -> dict:
        """Delete a document and its vectors from ChromaDB."""
        doc = session.get(Document, doc_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        # Delete from ChromaDB
        await delete_document_from_chroma(doc_id)

        # Delete file from disk
        if doc.storage_path and os.path.exists(doc.storage_path):
            try:
                os.remove(doc.storage_path)
            except Exception as e:
                print(f"Warning: Could not delete file {doc.storage_path}: {e}")

        # Delete from database
        session.delete(doc)
        session.commit()

        return {"message": f"Document '{doc.filename}' deleted successfully"}